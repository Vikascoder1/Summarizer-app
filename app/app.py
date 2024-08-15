# # app.py

# import streamlit as st
# from transformers import pipeline
# import pdfplumber
# import docx

# def read_pdf(file):
#     with pdfplumber.open(file) as pdf:
#         text = ''
#         for page in pdf.pages:
#             text += page.extract_text()
#     return text

# def read_docx(file):
#     doc = docx.Document(file)
#     text = [p.text for p in doc.paragraphs]
#     return '\n'.join(text)

# def summarize_text(text, model_name='facebook/bart-large-cnn', max_length=130, min_length=30):
#     summarizer = pipeline('summarization', model=model_name)
#     return summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)

# def main():
#     st.title("Document Summarizer")
#     st.write("Upload a document (PDF, DOCX, or TXT) to receive a summarized version.")

#     uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx', 'txt'])

#     if uploaded_file is not None:
#         if uploaded_file.type == 'application/pdf':
#             content = read_pdf(uploaded_file)
#         elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
#             content = read_docx(uploaded_file)
#         else:
#             content = uploaded_file.read().decode('utf-8')

#         st.write("Original Document:")
#         st.text_area("Content", content, height=300)

#         if st.button("Summarize"):
#             with st.spinner('Summarizing...'):
#                 summary = summarize_text(content)
#                 st.success("Summarization completed!")
#                 st.write("Summary:")
#                 st.text_area("Summary", summary[0]['summary_text'], height=200)

# if __name__ == '__main__':
#     main()

# app.py

# import streamlit as st
# from transformers import pipeline
# import pdfplumber
# import docx

# # Load the summarization model once
# summarizer = pipeline('summarization', model='facebook/bart-large-cnn')

# # Define helper functions
# def read_pdf(file):
#     text = []
#     with pdfplumber.open(file) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text.append(page_text)
#     return '\n'.join(text)

# def read_docx(file):
#     doc = docx.Document(file)
#     # Extract text from each paragraph and join with newline
#     text = [p.text for p in doc.paragraphs if p.text.strip()]
#     return '\n'.join(text)

# def summarize_text(text, max_length=130, min_length=30):
#     # Summarize the text using the pre-loaded summarizer
#     return summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)

# def main():
#     st.title("Document Summarizer")
#     st.write("Upload a document (PDF, DOCX, or TXT) to receive a summarized version.")

#     uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx', 'txt'])

#     if uploaded_file is not None:
#         # Read content based on file type
#         if uploaded_file.type == 'application/pdf':
#             content = read_pdf(uploaded_file)
#         elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
#             content = read_docx(uploaded_file)
#         elif uploaded_file.type == 'text/plain':
#             content = uploaded_file.read().decode('utf-8')
#         else:
#             st.error("Unsupported file type.")
#             return

#         # Display original document content
#         st.write("Original Document:")
#         st.text_area("Content", content, height=300)

#         # Summarize the content when the button is pressed
#         if st.button("Summarize"):
#             with st.spinner('Summarizing...'):
#                 try:
#                     summary = summarize_text(content)
#                     st.success("Summarization completed!")
#                     st.write("Summary:")
#                     st.text_area("Summary", summary[0]['summary_text'], height=200)
#                 except Exception as e:
#                     st.error(f"An error occurred during summarization: {e}")

# if __name__ == '__main__':
#     main()

import streamlit as st
from transformers import pipeline
import pdfplumber
import docx

# Load the summarization model once
@st.cache_resource
def load_summarizer():
    return pipeline('summarization', model='facebook/bart-large-cnn')

summarizer = load_summarizer()

# Define helper functions
def read_pdf(file):
    try:
        text = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n'.join(text)
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return ""

def read_docx(file):
    try:
        doc = docx.Document(file)
        text = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(text)
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return ""

def summarize_text(text, max_length=130, min_length=30):
    try:
        return summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)
    except Exception as e:
        st.error(f"Error summarizing text: {e}")
        return [{"summary_text": "Error generating summary"}]

def main():
    st.title("Document Summarizer")
    st.write("Upload a document (PDF, DOCX, or TXT) to receive a summarized version.")

    uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'docx', 'txt'])

    if uploaded_file is not None:
        # Read content based on file type
        if uploaded_file.type == 'application/pdf':
            content = read_pdf(uploaded_file)
        elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            content = read_docx(uploaded_file)
        elif uploaded_file.type == 'text/plain':
            content = uploaded_file.read().decode('utf-8')
        else:
            st.error("Unsupported file type.")
            return

        # Display original document content
        if content:
            st.write("Original Document:")
            st.text_area("Content", content, height=300)

            # Summarize the content when the button is pressed
            if st.button("Summarize"):
                with st.spinner('Summarizing...'):
                    summary = summarize_text(content)
                    if summary:
                        st.success("Summarization completed!")
                        st.write("Summary:")
                        st.text_area("Summary", summary[0]['summary_text'], height=200)

if __name__ == '__main__':
    main()
